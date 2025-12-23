from copy import deepcopy
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor, Pt, Cm
from docxcompose.composer import Composer
from docxtpl import DocxTemplate, Subdoc
import time
from mcp_bus.monkey_patch_map_tree import apply_fast_render_patch
import re
from utils.utils import to_chinese_num

apply_fast_render_patch()

from docx.oxml.ns import qn
from docx.oxml import parse_xml, OxmlElement
from lxml import etree

class Context:
    def __init__(self, app_session, sd):
        self.app_session = app_session
        self.docx = sd
        self.expand_data = {}

    def extract_element(self, source_path: str, xpath) -> Document:
        ele_index = -1
        ele_type = ''
        match = re.search(r'/w:([^[\]/]+)(?:\[[^\]]*\])?$', xpath)
        if match:
            ele_type = match.group(1)

        match = re.search(r'\[(\d+)\]', xpath)
        if match:
            ele_index = int(match.group(1))

        if ele_type == '' or ele_index < 0 :
            return None

        doc = Document(f"{self.app_session.app_dir}/{self.app_session.appctx_id}/{source_path}")
        body = doc.element.body
        ele_counter = 0
        for child in list(body):
            tag = child.tag
            if tag.endswith("sectPr"):
                continue
            if  tag.endswith(ele_type):
                ele_counter += 1
                if ele_counter == ele_index:
                    continue
            body.remove(child)
        return doc



    def _insert_element_as_paragraphs(self, elem, level,index):
        xml_str = etree.tostring(elem, encoding='unicode', with_tail=False)

        # 清除 numbering 相关的 w:numPr
        temp = parse_xml(xml_str)
        for numPr in temp.xpath('.//w:numPr'):
            numPr.getparent().remove(numPr)

        # 清除目录相关的 w:outlineLvl
        for outlineLvl in temp.xpath('.//w:outlineLvl'):
            outlineLvl.getparent().remove(outlineLvl)

        # 清除目录相关的 w:pStyle
        for pStyle in temp.xpath('.//w:pStyle'):
            pStyle.getparent().remove(pStyle)

        self.docx.element.body.insert(index,temp)

    def handle_xpath(self, node, index):
        source_doc = Document(self.source_file)
        if not source_doc:
            return False

        xpath = node.get("xpath")
        elements = source_doc._element.body.xpath(xpath)
        if not elements:
            self.app_session.write_log(f"未找到匹配 xpath: {xpath}")
            return False

        level = node.get('level', -1)
        if level == -1:
            for elem in elements:
                self._insert_element_as_paragraphs(elem, level,index)

        contents = node.get("content")
        if not isinstance(contents, list):
            return True

        handled_xpaths = []
        for content in contents:
            if content.get("xpath") in handled_xpaths:
                continue
            handled_xpaths.append(content.get("xpath"))
            if content["type"] in ("image", "table_image"):
                    doc_ele = self.extract_element(node.get("file_name"), content.get("xpath"))
                    if not doc_ele is None:
                        Composer(self.docx).insert(index, doc_ele)
                        index += 1
            else:
                self.handle_xpath(content, index)
                index += 1

        return True

    def apply_table_global_style(self, table):
        for row_idx, row in enumerate(table.rows):
            row.height = Cm(0.6)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for col_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = 1  # WD_CELL_VERTICAL_ALIGNMENT.CENTER = 1
                grid_span = cell._tc.grid_span
                for p in cell.paragraphs:
                    if grid_span > 1 :
                        p.alignment = 0
                    if col_idx == 0 or row_idx < self.expand_data.get("merge_row_idx", 1):
                        p.alignment = 1
                    p_pr = p._p.get_or_add_pPr()
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:lineRule'), 'exact')  # 固定值模式（exact）
                    spacing.set(qn('w:line'), '240')  # 12 磅 = 240 缇（1 磅 = 20 缇）
                    p_pr.append(spacing)

                    for child in list(p_pr):
                        if child.tag == qn('w:pStyle'):
                            p_pr.remove(child)

                    for r in p.runs:
                        r.font.name = '宋体'
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0, 0, 0)
                        r.bold = True if grid_span > 1 or row_idx < self.expand_data.get("merge_row_idx", 1) else False

    def set_cell_text(self, tr, cell_index, text):
        tc = tr.find(f'.//w:tc[{cell_index}]', namespaces=tr.nsmap)
        for p in tc.findall('.//w:p', namespaces=tr.nsmap):
            p.getparent().remove(p)
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = str(text)
        r.append(t)
        p.append(r)
        tc.append(p)

    def find_max_row_index_by_group_text(self, table, by_text):
        found = False
        group_row_index = -1
        group_data_index = 0
        for row_index, row in enumerate(table.rows):
            first_cell_text = row.cells[0].text.strip()
            second_cell_text = row.cells[1].text.strip()
            if second_cell_text == by_text:
                found = True
                group_row_index = row_index + 1
                continue
            if found:
                if first_cell_text.isdigit():
                    group_row_index = row_index + 1
                    group_data_index = int(first_cell_text)
                else:
                    break
        return group_row_index, group_data_index

    def merge_table(self, target_tbl, source_tbl, group):
        if target_tbl is None:
            header_tbl = deepcopy(source_tbl)._tbl
            self.expand_data["merge_row_idx"] = 0
            is_header = False
            for i, tr in enumerate(header_tbl.xpath('.//w:tr')):
                if i ==0:
                    is_header = True
                    self.expand_data["merge_row_idx"] += 1
                    continue
                if is_header and (
                        (len(tr.xpath('.//w:vMerge[@w:val="restart"]')) > 0) or
                        (len(tr.xpath('.//w:vMerge[@w:val="continue"]')) > 0)
                ):
                    self.expand_data["merge_row_idx"] += 1
                    continue
                is_header = False
                header_tbl.remove(tr)

            self.docx.element.body.append(header_tbl)
            target_tbl= self.docx.tables[-1]

            self.expand_data["row_no"] = 0
            self.expand_data["sub_row_no"] = 0
            self.expand_data["group_row_index"] = 0
            self.expand_data["group_data_index"] = 0

        if group == "ITEM-TYPE":
            for row_index, row in enumerate(source_tbl.rows[self.expand_data["merge_row_idx"]:],start = self.expand_data["merge_row_idx"]):
                first_cell_text = row.cells[0].text.strip()
                second_cell_text = row.cells[1].text.strip()
                if second_cell_text.strip() == '':
                    continue
                if first_cell_text.isdigit():
                    data_rows = source_tbl._tbl.xpath(f'./w:tr[position()={row_index + 1}]')
                    for tr in data_rows:
                        self.expand_data["group_row_index"] += 1
                        self.expand_data["group_data_index"] += 1
                        self.set_cell_text(tr, 1, self.expand_data["group_data_index"])
                        target_tbl._tbl.insert(self.expand_data["group_row_index"], deepcopy(tr))
                else:
                    group_row_index, group_data_index = self.find_max_row_index_by_group_text(target_tbl, second_cell_text)
                    if group_row_index > -1:
                        self.expand_data["group_row_index"] = group_row_index + 1
                        self.expand_data["group_data_index"] = group_data_index
                    else:
                        new_row = target_tbl.add_row()
                        first_cell = new_row.cells[1]
                        first_tc = first_cell._tc
                        tcPr = first_tc.get_or_add_tcPr()
                        gridSpan = OxmlElement('w:gridSpan')
                        gridSpan.set(qn('w:val'), '4')
                        tcPr.append(gridSpan)

                        tr = new_row._tr
                        tcs = tr.findall('.//w:tc', namespaces=tr.nsmap)
                        for tc in tcs[2: 5]:
                            tr.remove(tc)

                        self.expand_data["row_no"] += 1
                        self.expand_data["group_row_index"] = len(target_tbl.rows) + 1
                        self.expand_data["group_data_index"] = 0

                        self.set_cell_text(tr, 1, to_chinese_num(self.expand_data["row_no"]))
                        self.set_cell_text(tr, 2, second_cell_text)
        else:
            if group == "ORG":
                # 增加单位名称行 合并序号后面单元格
                new_row = target_tbl.add_row()
                first_cell = new_row.cells[1]
                first_tc = first_cell._tc
                tcPr = first_tc.get_or_add_tcPr()
                gridSpan = OxmlElement('w:gridSpan')
                gridSpan.set(qn('w:val'), '4')
                tcPr.append(gridSpan)

                tr = new_row._tr
                tcs = tr.findall('.//w:tc', namespaces=tr.nsmap)
                for tc in tcs[2: 5]:
                    tr.remove(tc)

                self.expand_data["row_no"] += 1

                self.set_cell_text(tr, 1, to_chinese_num(self.expand_data["row_no"]))
                self.set_cell_text(tr, 2, self.company)

            # 拷贝数据行
            data_rows = source_tbl._tbl.xpath(f'./w:tr[position()>{self.expand_data["merge_row_idx"]}]')
            for tr in data_rows:
                cells = tr.xpath('./w:tc')
                if len(cells) >= 2:
                    second_cell = cells[1]
                    texts = second_cell.xpath('.//w:t')
                    text = ''.join(t.text or '' for t in texts)
                    if text.strip() == '':
                        continue

                self.expand_data["sub_row_no"] += 1
                self.set_cell_text(tr, 1, self.expand_data["sub_row_no"])
                target_tbl._tbl.append(deepcopy(tr))

        self.apply_table_global_style(target_tbl)

    def merge_table_with_xpath(self, node, table_index = -1, group = None):
        source_doc = Document(self.source_file)
        if not source_doc:
            return False

        xpath = node.get("xpath")
        elements = source_doc._element.body.xpath(xpath)
        if not elements:
            self.app_session.write_log(f"未找到匹配 xpath: {xpath}")
            return False

        contents = node.get("content")
        if not isinstance(contents, list):
            return True

        tbl = self.docx.tables[-1] if self.docx.tables else None
        content_table_index = 0
        for content in contents:
            if not content["type"] in ("table"):
                continue
            else:
                content_table_index+=1

            if table_index >= -1 and table_index != content_table_index:
                continue
            match = re.search(r'\[(\d+)\]', content.get("xpath"))
            if match:
                table_index = int(match.group(1)) -1
                self.merge_table(tbl, source_doc.tables[table_index], group)


    def handle(self, item, action, table_index = -1, group = None):
        if item is None:
            return
        if 'level' not in item:
            return

        if "file_name" in item:
            self.source_file =  self.app_session.app_dir / self.app_session.appctx_id / item["file_name"]

        if "company" in item:
            self.company =  item["company"]

        if action == "copy":
            level = item["level"]
            title = item["title"]
            header = self.docx.add_heading(title, level=level + 1)
            index = self.docx._element.body.index(header._element)

            xpath = item["xpath"]
            if xpath:
                self.handle_xpath(item,index + 1)
            else:
                content = item["content"]
                if isinstance(content, list):
                    content = "\n\n".join(item.get("content", "") for item in content)
                self.docx.add_paragraph(content)
        elif action == "merge":
            self.merge_table_with_xpath(item, table_index, group)


def traverse_tree(ctx, parent, node, level, action = None, table_index = -1, group = None):
    ctx.app_session.write_log(f"node={node}")
    if action is None:
        action = node.get("action", None)
        table_index = node.get("table_index", -1)
        group = node.get("group", None)

    if node.get("level", 0) >= level:
        ctx.handle(node, action, table_index, group)

    rels = ctx.app_session.graph.get_relationship(
        {
            "src_label": "Data",
            "src_props": {"node_id": node['node_id']},
            "rel_types": ['SUBHEADING'],
            "hop_num": 1
        }
    )

    rels = sorted(rels, key=lambda odr: odr[2].get("order", 0))

    for _, _, n in rels:
        traverse_tree(ctx, node, n, level, action, table_index, group)

def write_tree(app_session, tpl, var_node):
    sd = tpl.new_subdoc()
    ctx = Context(app_session, sd)
    traverse_tree(ctx, None, var_node, 0)
    return sd

def find_node_by_name(app_session, variable_name):
    var_node = app_session.graph.get_node("Data", {
        "app_ctx_id": app_session.appctx_id,
        "name": variable_name
    })
    if not var_node:
        return None
    return var_node[0]

def write_tmpl(app_session, template_file_name, filename):
    tmpl = app_session.app_dir / "const" / template_file_name
    if not tmpl.exists():
        return False, "template file not found"
    tpl = DocxTemplate(tmpl)

    variable_names = tpl.get_undeclared_template_variables()

    variables = {}
    variables_no_data = []
    for variable_name in variable_names:
        variable_node = find_node_by_name(app_session, variable_name)
        if variable_node is None:
            variables_no_data.append(variable_name)
        else:
            variables[variable_name] = variable_node

    if variables_no_data:
        return False, '\n\n'.join(f"{{{var}}} - 数据未就绪" for var in variables_no_data)

    context = {
        variable_name: write_tree(app_session, tpl, variables[variable_name])
        for variable_name, node in variables.items()
    }

    t = time.time()
    tpl.render(context)
    print(f"elapse time: {time.time() - t}")

    doc_out_dir = app_session.app_dir / app_session.appctx_id
    doc_out_dir.mkdir(parents=True, exist_ok=True)
    if not filename.endswith(".docx"):
        filename = f"{filename}.docx"
    doc_out_name = doc_out_dir / f"{filename}"
    tpl.save(doc_out_name)
    return True, f"{app_session.app_id}/{app_session.appctx_id}/{filename}"


if __name__ == '__main__':
    output = r"/home/woody/seektime/PeritusLLM/is_kb/HZ/output/2025年8月工程信息月报（工程管理部）.docx"
    write_templ()